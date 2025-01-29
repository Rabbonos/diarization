from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import List, Optional
import datetime
import os
import shutil
from typing import Callable
from settings import Settings

# для меток
def format_time_range(start, end):
        def format_time(seconds):
            td = datetime.timedelta(seconds=seconds)
            # Get full seconds and remaining milliseconds
            full_seconds = int(td.total_seconds())
            milliseconds = int((td.total_seconds() - full_seconds) * 1000)
            # Format time in hours:minutes:seconds,milliseconds
            formatted_time = str(datetime.timedelta(seconds=full_seconds)) + f",{milliseconds:03d}"
            return formatted_time

        start_time = format_time(start)
        end_time = format_time(end)
        return f"[{start_time} --> {end_time}]"

def format_word(document:Callable , word_line_part1:str, word_line_part2:str)->None:
        '''
        Formats word-document's body
        '''

        line_runs = document.add_paragraph(word_line_part1)
        line_run=line_runs.runs[0]
        line_run.font.name = "Calibri"
        line_run.font.size = Pt(11)
        line_run.bold = True 

        line_run = line_runs.add_run(word_line_part2)
        line_run.bold=False

def merge_neighbor_segments(segments:List[dict], assigned_speakers:List[str], settings:Settings):
       """
        Merges consecutive segments if speakers are the same.
        
        Args:
            segments (List[dict]): A list of segment dictionaries, each containing "start" and "text" keys.
            assigned_speakers (List[str]): A list of speakers corresponding to each segment.
            
        Returns:
            list: A list of merged segments.
       """
       merged_segments  = [segments[0]]
       for i in range(1,len(segments)):
                previous_speaker= assigned_speakers[i-1]
                current_speaker = assigned_speakers[i]

                if previous_speaker == current_speaker and current_speaker != settings.undefiend_speaker:
                       merged_segments[-1]['end'] =segments[i]["end"]
                       merged_segments[-1]['text'] = segments[i-1]['text'] + segments[i]['text']                   
                else:
                       merged_segments.append(segments[i])
 
       return merged_segments


def write_to_word(segments:List[dict], assigned_speakers:List[str], duration, settings, include_timestamps:Optional[bool]=True,  output_folder:str=None):
        """
        Writes transcription data to a Word file.

        Args:
            segments (List[dict]): A list of segment dictionaries, each containing "start" and "text" keys.
            assigned_speakers (List[str]): A list of speakers corresponding to each segment.
            duration (int): Total duration of the audio in seconds.
            output_path (str): Path to save the Word document. Defaults to 'transcription.docx'.
            include_timestamps (bool): If True, includes timestamps in the transcription.

        Returns:
            None
        """
        segments = merge_neighbor_segments(segments, assigned_speakers, settings)

        if not output_folder:
               output_folder = os.getcwd()
               
        now = datetime.datetime.now().strftime("%Y-%m-%d %H-%M-%S")
        title = f"Расшифровка итоговая от {now}"
        output_path= os.path.join(output_folder,title)

        # Create a new Word document
        document = Document()
        title = document.add_paragraph(title)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title.runs[0]
        title_run.font.name = "Calibri"
        title_run.font.size = Pt(16)
        title_run.bold = True 

        for i in range(len(segments)):
                speaker_name = assigned_speakers[i]
                segment = segments[i]
                end_time = segments[i + 1]["start"] if i + 1 < len(segments) else duration
                time_range = format_time_range(segment["start"], end_time)
                word_line_part2=f"{segment['text']}"
                if include_timestamps:
                        word_line_part1=f"\n{time_range} {speaker_name}:"
                        format_word(document, word_line_part1, word_line_part2)
                else:
                        word_line_part1=f"\n{speaker_name}:"
                        format_word(document, word_line_part1, word_line_part2)

        for paragraph in document.paragraphs:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)

        # Save the document
        document.save(output_path)

        print(f"Word file saved as {output_path}")

            
def save_transcription(file_path:str, segments:List[dict], assigned_speakers:List[str], duration:int, include_timestamps:Optional[bool]=True)->None:
        """
        Saves transcription to a txt file.

        Args:
            file_path (str): Path to save the transcription.
            segments (List[dict]): List of segment dictionaries containing text and timestamps.
            assigned_speakers (List[str]): List of speaker names assigned to each segment.
            duration (int): Duration of the entire audio in seconds.
            save_mode(str):how to save the file
            include_timestamps (bool, optional): If True, includes time markers in the transcription. Defaults to True.

        Returns:
            None
        """
        
        # Write the transcript with formatted time ranges
        with open(file_path, "w", encoding='utf-8') as f:
            for i in range(len(segments)):
                speaker_name = assigned_speakers[i]
                segment = segments[i]
                end_time = segments[i + 1]["start"] if i + 1 < len(segments) else duration
                time_range = format_time_range(segment["start"], end_time)
                if include_timestamps:
                        f.write(f"\n{time_range} {speaker_name}: {segment['text']}")
                else:
                        f.write(f"\n{speaker_name}: {segment['text']}")

        print('\nсохранено в "' + file_path + '"')


def print_transcription(segments:List[dict], assigned_speakers:List[str], duration:int, include_timestamps:Optional[bool]=True)->None:
     """
     Writes transcription data to a Word file.

     Args:
        segments (List[dict]): List of segment dictionaries containing text and timestamps.
        assigned_speakers (List[str]): List of speaker names assigned to each segment.
        duration (int): Duration of the entire audio in seconds.
        metki (bool, optional): If True, includes time markers in the transcription. Defaults to True.

     Returns:
        None
     """
     if include_timestamps:
            for i in range(len(segments)):
                            speaker_name = assigned_speakers[i]
                            segment = segments[i]
                            end_time = segments[i + 1]["start"] if i + 1 < len(segments) else duration
                            time_range = format_time_range(segment["start"], end_time)
                            if include_timestamps:
                                    print(f"\n{time_range} {speaker_name}: {segment['text']}")
                            else:
                                    print(f"\n{speaker_name}: {segment['text']}")


def remove_junk(path: str):
            """
            Removes a file or directory.

            Args:
                path (str): The file or directory path to remove.
            
            Raises:
                FileNotFoundError: If the path does not exist.
                OSError: If an error occurs during removal.
            """
            if os.path.isfile(path) or os.path.islink(path):  # Check for file or symlink
                os.remove(path)
            elif os.path.isdir(path):  # Check for directory
                shutil.rmtree(path)
            else:
                raise FileNotFoundError(f"The path {path} does not exist.")