#https://github.com/zedr/clean-code-python

# # !pip install nemo_toolkit[all] (при небходимости раздокументировать)
# !pip install onnxruntime
# !pip install pyannote-whisper
# !pip install yadisk[sync_defaults]
# !pip install -q gdown
# !pip install setuptools==59.5.0
# !pip install speechbrain==0.5.16
# !git clone https://github.com/yinruiqing/pyannote-whisper.git
# !cd pyannote-whisper && pip install .

#use rclone to mount google drive?


#TO ADD: IF only someone did not have sample, only him to denote as 'undefinde', not all 
#TO ADD: tones 
#TO ADD: TRAINING 
#TO ADD: TRAINING DATA PREPARATION
#TO ADD: async stuff + caching


import datetime
import subprocess
from pyannote.audio import Audio
import shutil
from pyannote.core import Segment
import torchaudio
from speechbrain.inference.separation import SepformerSeparation
import wave
import contextlib
import numpy as np
import os
import warnings
from typing import List, Dict, Callable , Tuple , Optional
#from google.colab import files
import re
from sklearn.metrics import pairwise_distances
#from google.colab import userdata
# import nemo.collections.asr as nemo_asr (раздокументировать при необходимости)
import soundfile as sf
import onnxruntime
import json
import yadisk #yandex

from audio_processing import fragment, download_audio, preprocess_segments, main_audio_preprocessing, AudioProcessor, get_sample_embeddings , replace_overlaps, use_vectors , get_embeddings_main_audio
from clustering import ClusteringModel
from file_manager import FolderTree, fetch_vectors_and_audio_files
from constants import *
from video import get_duration



def read_json(json_path):
            with open(json_path, 'r', encoding='utf-8') as f:
                return json.load(f)

def get_new_ytoken():
        
        '''get new yandex token'''

        with yadisk.Client(client_id, client_secret) as client:
            url = client.get_code_url()

            print(f"Go to the following url: {url}")
            code = input("Enter the confirmation code: ")

            try:
                response = client.get_token(code)
            except yadisk.exceptions.BadRequestError:
                print("Bad code")
                return
            
            client.token = response.access_token
            token_yandex=client.token
            print(client.token)
            if client.check_token():

                print("Sucessfully received token!")
            else:
                print("Something went wrong. Not sure how though...")

# для меток
def format_time_range(start, end):
        def format_time(seconds):
            td = datetime.timedelta(seconds=seconds)
            # Get full seconds and remaining milliseconds
            full_seconds = int(td.total_seconds())
            milliseconds = int((td.total_seconds() - full_seconds) * 1000)
            # Format time in hours:minutes:seconds,milliseconds
            formatted_time = str(datetime.timedelta(seconds=full_seconds)) + f",{milliseconds:03d}"
            return formatted_time

        start_time = format_time(start)
        end_time = format_time(end)
        return f"[{start_time} --> {end_time}]"



def remove_junk(path: str):
            """
            Removes a file or directory.

            Args:
                path (str): The file or directory path to remove.
            
            Raises:
                FileNotFoundError: If the path does not exist.
                OSError: If an error occurs during removal.
            """
            if os.path.isfile(path) or os.path.islink(path):  # Check for file or symlink
                os.remove(path)
            elif os.path.isdir(path):  # Check for directory
                shutil.rmtree(path)
            else:
                raise FileNotFoundError(f"The path {path} does not exist.")



def assign_speakers(clustering_manager:ClusteringModel, embeddings:List[np.array], segments:List[dict] , sample_embeddings:List[Tuple], SIMILARITY_THRESHOLD:float=None, UNASSIGNED_LABEL = "Участник (не определён)" ) -> List[str]:
        '''
        Assigns speakers by comparing embeddings of audio segments with samples.
    
        Args:
            clustering_manager (ClusteringModel): Object that handles clustering.
            embeddings (np.ndarray): 2D array of shape (n_segments, embedding_dim) for segment embeddings.
            segments (List[dict]): List of dictionaries describing audio segments.
            sample_embeddings (List[Tuple]): List of (embedding, speaker_name) tuples for known speakers.
            similarity_threshold (float, optional): Threshold for cosine similarity. Defaults to None.

        Returns:
            List[str]: Speaker names or `UNASSIGNED_LABEL` for each segment.

        We start with list of embeddings (for each segment)
        We then cluster the embeddings of the segments, grouping similar segments together. 
        Once the segments are clustered, we assign a speaker label to each one by comparing the cluster’s centroid with known speaker samples.
        '''
        cluster_labels= clustering_manager.cluster(embeddings)###
        assigned_speakers = [None] * len(segments)
        cluster_to_speaker = {}

        for cluster_id in np.unique(cluster_labels):

            cluster_indices = np.where(cluster_labels == cluster_id)[0]
            cluster_embeddings = embeddings[cluster_indices]
            centroid = np.mean(cluster_embeddings, axis=0)

            #norm - (p-norm) , default is p-2 (L2) norm a.k.a euclidian distance
            similarities = [
                (np.dot(centroid, sample_embedding.T) / (np.linalg.norm(centroid) * np.linalg.norm(sample_embedding)), person_name) #cosine similarity
                for sample_embedding, person_name in sample_embeddings
                            ]
            
            similarity_scores = [similarity for similarity, person_name in similarities]
            if SIMILARITY_THRESHOLD == None:
                mean_similarity = np.mean(similarity_scores)
                std_similarity = np.std(similarity_scores)
                SIMILARITY_THRESHOLD = mean_similarity + std_similarity/1 ############################

            max_similarity, most_similar_person = max(similarities, key=lambda x: x[0])

            if max_similarity >= SIMILARITY_THRESHOLD:
                cluster_to_speaker[cluster_id] = most_similar_person
            else:
                cluster_to_speaker[cluster_id] = UNASSIGNED_LABEL 

        for i, segment in enumerate(segments):
            if cluster_labels[i] in cluster_to_speaker:
                 assigned_speakers[i] = cluster_to_speaker[cluster_labels[i]]
            else:
                 warnings.warn(f"Segment {i} with cluster ID {cluster_labels[i]} was not assigned a speaker.")

        return assigned_speakers

def assign_speakers_individually(segments:List[dict], embeddings :List[np.ndarray],  sample_embeddings: List[Tuple[np.ndarray, str]], SIMILARITY_THRESHOLD: Optional[float] = None, UNASSIGNED_LABEL = "Участник (не определён)" ) -> List[str]:
            """
            Assign speakers to audio segments by comparing each segment embedding with all sample embeddings.
            Similar to function 'assign_speakers', main difference - no clustering done here.

            Args:
                segments (List[dict]): List of audio segments.
                embeddings (List[np.ndarray]): List of embeddings for each segment.
                sample_embeddings (List[Tuple[np.ndarray, str]]): List of tuples containing sample embeddings and corresponding speaker names.
                SIMILARITY_THRESHOLD (float, optional): The threshold for similarity to consider a segment as matching a speaker. 
                                                    If not provided, it will be calculated as the mean similarity plus one standard deviation.

            Returns:
                List[str]: List of speaker names assigned to each segment.
            """

            assigned_speakers = [''] * len(segments)
            
            for i, embedding in enumerate(embeddings):

                similarities = [
                        (np.dot(embedding, sample_embedding.T) / (np.linalg.norm(embedding) * np.linalg.norm(sample_embedding)), person_name) #cosine similarity
                        for sample_embedding, person_name in sample_embeddings
                                ]
                similarity_scores = [similarity for similarity, _ in similarities]
                max_similarity, most_similar_person = max(similarities, key=lambda x: x[0])

                if SIMILARITY_THRESHOLD == None:
                    mean_similarity = np.mean(similarity_scores)
                    std_similarity = np.std(similarity_scores)
                    SIMILARITY_THRESHOLD = mean_similarity + std_similarity/1

                if max_similarity > SIMILARITY_THRESHOLD:
                    assigned_speakers[i] = most_similar_person
                else:
                    assigned_speakers[i] = UNASSIGNED_LABEL
                    #No need now, will need if add numbers to unassigned speakers
                    #sample_embeddings.append([ embedding , UNASSIGNED_LABEL]) #Adding new embedding, when useful: to group all unassigned similar speakers together
                    warnings.warn(f"Segment {i} not assigned to any known speaker.")
                
            return assign_speakers


from docx import Document

def write_to_word(segments:List[dict], assigned_speakers:List[str], duration, include_timestamps:Optional[bool]=True,  output_path:str='transcription.docx'):
        """
        Writes transcription data to a Word file.

        Args:
            segments (List[dict]): A list of segment dictionaries, each containing "start" and "text" keys.
            assigned_speakers (List[str]): A list of speakers corresponding to each segment.
            duration (int): Total duration of the audio in seconds.
            output_path (str): Path to save the Word document. Defaults to 'transcription.docx'.
            include_timestamps (bool): If True, includes timestamps in the transcription.

        Returns:
            None
        """
        # Create a new Word document
        document = Document()

        for i in range(len(segments)):
                speaker_name = assigned_speakers[i]
                segment = segments[i]
                end_time = segments[i + 1]["start"] if i + 1 < len(segments) else duration
                time_range = format_time_range(segment["start"], end_time)
                if include_timestamps:
                        document.add_paragraph(f"\n{time_range} {speaker_name}: {segment['text']}")
                else:
                        document.add_paragraph(f"\n{speaker_name}: {segment['text']}")

        # Save the document
        document.save(output_path)

        print(f"Word file saved as {output_path}")

            
def save_transcription(file_path:str, segments:List[dict], assigned_speakers:List[str], duration:int, include_timestamps:Optional[bool]=True)->None:
        """
        Saves transcription to a txt file.

        Args:
            file_path (str): Path to save the transcription.
            segments (List[dict]): List of segment dictionaries containing text and timestamps.
            assigned_speakers (List[str]): List of speaker names assigned to each segment.
            duration (int): Duration of the entire audio in seconds.
            save_mode(str):how to save the file
            include_timestamps (bool, optional): If True, includes time markers in the transcription. Defaults to True.

        Returns:
            None
        """
        
        # Write the transcript with formatted time ranges
        with open(file_path, "w", encoding='utf-8') as f:
            for i in range(len(segments)):
                speaker_name = assigned_speakers[i]
                segment = segments[i]
                end_time = segments[i + 1]["start"] if i + 1 < len(segments) else duration
                time_range = format_time_range(segment["start"], end_time)
                if include_timestamps:
                        f.write(f"\n{time_range} {speaker_name}: {segment['text']}")
                else:
                        f.write(f"\n{speaker_name}: {segment['text']}")

        print('\nсохранено в "' + file_path + '"')


def print_transcription(segments:List[dict], assigned_speakers:List[str], duration:int, include_timestamps:Optional[bool]=True)->None:
     """
     Writes transcription data to a Word file.

     Args:
        segments (List[dict]): List of segment dictionaries containing text and timestamps.
        assigned_speakers (List[str]): List of speaker names assigned to each segment.
        duration (int): Duration of the entire audio in seconds.
        metki (bool, optional): If True, includes time markers in the transcription. Defaults to True.

     Returns:
        None
     """
     if include_timestamps:
            for i in range(len(segments)):
                            speaker_name = assigned_speakers[i]
                            segment = segments[i]
                            end_time = segments[i + 1]["start"] if i + 1 < len(segments) else duration
                            time_range = format_time_range(segment["start"], end_time)
                            if include_timestamps:
                                    print(f"\n{time_range} {speaker_name}: {segment['text']}")
                            else:
                                    print(f"\n{speaker_name}: {segment['text']}")

def start():
            global Accuracy_boost ,huggingface_token, main_folder , main_audio, participants, folders, language, modeltype, embedding_model_name, Voice_sample_exists, Vector, Add, save_txt, metki, SIMILARITY_THRESHOLD, n_clusters, DISTANCE_THRESHOLD, clustering, settings, Silence
            global folders_restore, participants_restore , get_token, token_yandex , use_yandex, client , remove_overlap

            main_audio_restore= main_audio
            participants_restore = participants
            use_yandex=False
            speaker_model=None 
            junk = []
            junk.append(TEMP_FOLDER_MAIN, TEMP_FOLDER_FRAGMENTS, CLIPPED_SEGMENTS) #for now only this

            #используем яндекс диск
            if mode=='yandex':
                use_yandex=True
                client = yadisk.Client(token=token_yandex  )
            else:
                 client=None

            if get_token =='да':
                   get_new_ytoken()

            #главный файл (путь)
            #main auido - ssilka / abs. path k audio
            # if file does not exist in local system
            if not os.path.isfile(main_audio):              
                       main_audio =  download_audio(source, main_audio)

            print(main_audio, 'MAIN AUDIO FILE"S PATH')

            #folders / files , sample folders here
            folders_manager = FolderTree(main_folder, 'yandex', ATTEMPTS, ATTEMPTS_INTERVAL, client)

            participant_folder_paths , participants = folders_manager.process_subfolders()
            #participants restore
            #folders restore

            ###FRAGEMENTS HERE
            fragments = fragment(interval, main_audio , divide_interval, TEMP_FOLDER_MAIN, TEMP_FOLDER_FRAGMENTS )
            
            print(participant_folder_paths, 'participant_folder_paths')

            #HERE VECTORS AND AUDIO DATA
            sample_vectors, sample_audios =  fetch_vectors_and_audio_files(participant_folder_paths, participants, mode, client, ATTEMPTS, ATTEMPTS_INTERVAL)

            print(sample_vectors, 'sample_vectors', sample_audios, 'sample_audios')

            # Remove participants who don't have any audio samples
            participants_with_samples = [name for name in participants if any(name == name_ for _, name_, _ in sample_audios)]

            # Create a new list of folders corresponding to participants with samples
            folders_with_samples = [folder for folder, name in zip(participant_folder_paths, participants) if name in participants_with_samples]

            if len(participants_with_samples)!=len(participants):
                warnings.warn(f'Some participants did not have audio samples, listing those that had \
                        {participants_with_samples}')

            # Update the participants and folders lists
            participants = participants_with_samples
            folders = folders_with_samples
            print('participants_with_samples', participants_with_samples, folders_with_samples, ' folders_with_samples')

            #zdes u nas iz sample_vectors, sample_audios est participants bez audio
            sample_vectors, sample_audios =  fetch_vectors_and_audio_files(participant_folder_paths, participants, mode, client, ATTEMPTS, ATTEMPTS_INTERVAL)
            
            print('POSLE UDALENIYA BEZ AUDIO PARTICIPANTS:',sample_vectors, 'sample_vectors', sample_audios, 'sample_audios')

            # A LOT OF AUDIPROCESSING PREPARATIONS HERE
            audio_models_loading = AudioProcessor(language, modeltype, embedding_model_name, HUGGINGFACE_TOKEN, Accuracy_boost, Silence )
            embedding_model= AudioProcessor.embedding_model
            embedding_model_dimension = AudioProcessor.embedding_model_dimension
            model= AudioProcessor.model
            vad_pipeline=None 
            pipeline=None
            embedding_models_group1 = audio_models_loading.embedding_models_group1

            if Silence:         
                  vad_pipeline= AudioProcessor.vad_pipeline

            if Accuracy_boost:
                  pipeline= AudioProcessor.diarization_pipeline

            if save_txt:
                        file_path = f"Определение голосов расшифровки Whisper - {main_audio_restore}.txt"
            else:
                  file_path=None
            #основная функция
            def process_file(audio_path, sample_folders, model, metki, save_txt, n_clusters=None, SIMILARITY_THRESHOLD=None, DISTANCE_THRESHOLD =None , Vector=False, Add=False, Voice_sample_exists=False, clustering=True, Silence = False, Accuracy_boost=False):
                
                    #I already checked this beforehand, now need to delete this and put here variable I got from before 
                    names_in_sample_audios=[name for file,name,_ in sample_audios]
                    print (len(set(names_in_sample_audios)), len(participants), 'SRAVNENIYE DLIN DLYA VOICE SAMPLE EXISTS!')
                    if len(set(names_in_sample_audios))!= len(participants):
                            Voice_sample_exists=False

                    #main_audio loading 
                    result = main_audio_preprocessing(model,'audio.wav',is_fragment, main_audio, fragments, Silence, Accuracy_boost, vad_pipeline, pipeline)
                    duration = get_duration('audio.wav')
                    segments = result["segments"]
                    sep_model = SepformerSeparation.from_hparams( "speechbrain/sepformer-whamr16k" )
                    overlap_timestamps = AudioProcessor.remove_overlap()

                    # deleting old segments and replacing them with new segments HERE
                    segments = replace_overlaps(segments, sep_model, model, overlap_timestamps)
                    #sorting because I potentially changed the order in the previous lines
                    segments = sorted(segments, key=lambda segment: segment['start'])
                    segments = preprocess_segments(segments)

                    print(segments, 'SEGMENTS')

                    #GETTING SAMPLE VECTORS OR USING PRE-MADE SAMPLE VECTORS
                    if Vector:
                          sample_embeddings = use_vectors(sample_vectors, client)
                    else:
                         if Voice_sample_exists:
                                        sample_embeddings = get_sample_embeddings(sample_audios, sample_vectors, embedding_model, speaker_model, embedding_models_group1, client)

                    #GETTING EMBEDDINGS OF MAIN AUDIO
                    embeddings = get_embeddings_main_audio('audio.wav', segments, embedding_model, embedding_model_dimension, embedding_models_group1, speaker_model)
                    embeddings = np.nan_to_num(embeddings)
                    embedding_size = embeddings.shape[1]
                    small_vector = np.full(embedding_size, -1e10)  # or np.zeros(embedding_size)
                    # Replace rows where embeddings are all zeros with small_vector
                    embeddings[~np.any(embeddings, axis=1)] = small_vector


                    distances = pairwise_distances(embeddings, metric='cosine')  # or use any other metric
                    distances = distances[np.triu_indices_from(distances, k=1)]

                    if DISTANCE_THRESHOLD==None:
                          DISTANCE_THRESHOLD = np.mean(distances) +  np.std(distances)/2

                    #CLUSTERING BLOCK
                    clustering_manager = ClusteringModel(settings)

                    if Voice_sample_exists:
                      if clustering:
                            assigned_speakers = assign_speakers(clustering_manager, embeddings, segments, sample_embeddings, SIMILARITY_THRESHOLD)
                          
                    if not clustering:
                       assigned_speakers = assign_speakers_individually(segments, embeddings, sample_embeddings)

                    #MAKE IT DIFFERENT, READ 'TO ADD'
                    if not Voice_sample_exists:
                        cluster_labels= clustering_manager.cluster(embeddings) ###
                        assigned_speakers = ["Участник (не определён)" for i in cluster_labels]

                    #OUTPUT BLOCK
                    if save_mode == 'text':
                            save_transcription(file_path, segments, assigned_speakers, duration, metki)
                    elif save_mode == 'word':
                            write_to_word(segments,assigned_speakers , duration, metki, 'transcription.docx')
                    else:
                            print_transcription(segments, assigned_speakers, duration, metki)

            #??
            process_file(main_audio, sample_folders, model , metki, save_txt, n_clusters, SIMILARITY_THRESHOLD, DISTANCE_THRESHOLD, Vector, Add, Voice_sample_exists, clustering, Silence, Accuracy_boost) #запуск кода

            #проверяем результат, работает если метка save_txt = True
            if file_path!=None:
                try:
                    with open(file_path, "r", encoding='utf-8') as f:
                        transcript = f.read()
                    print(transcript)
                except Exception as e:
                    print(f'For unknown reasons transcription could not be printed, {str(e)}')

            remove_junk(junk)
              